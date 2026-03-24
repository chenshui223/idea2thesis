from __future__ import annotations

from pathlib import Path

from docx import Document

from idea2thesis.contracts import ParsedBrief


def _split_labeled_values(paragraph: str) -> list[str]:
    _, _, content = paragraph.partition("：")
    return [item.strip() for item in content.split("、") if item.strip()]


def parse_brief(file_path: Path) -> ParsedBrief:
    document = Document(file_path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    raw_text = "\n".join(paragraphs)

    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    title = paragraphs[0] if paragraphs else "Untitled Brief"
    requirements: list[str] = []
    constraints: list[str] = []
    tech_hints: list[str] = []
    thesis_cues: list[str] = []
    collect_tech_bullets = False

    for paragraph in document.paragraphs[1:]:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        is_bullet = "Bullet" in style_name
        if text.startswith("功能要求："):
            requirements.extend(_split_labeled_values(text))
            collect_tech_bullets = False
        elif text.startswith("约束条件："):
            constraints.extend(_split_labeled_values(text))
            collect_tech_bullets = False
        elif text.startswith("论文提纲："):
            thesis_cues.extend(_split_labeled_values(text))
            collect_tech_bullets = False
        elif text.startswith("技术要求"):
            collect_tech_bullets = True
        elif collect_tech_bullets and is_bullet:
            tech_hints.append(text)
        elif "Python" in text or "Java" in text or "React" in text:
            tech_hints.append(text)

    if tables:
        thesis_cues.extend(
            [row[0] for table in tables for row in table[1:] if row and row[0]]
        )

    return ParsedBrief(
        title=title,
        requirements=requirements,
        constraints=constraints,
        tech_hints=tech_hints,
        thesis_cues=list(dict.fromkeys(thesis_cues)),
        raw_text=raw_text,
        extraction_snapshot={"paragraphs": paragraphs, "tables": tables},
    )
