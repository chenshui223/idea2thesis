from pathlib import Path

from docx import Document

from idea2thesis.parser import parse_brief


def test_parse_brief_extracts_structured_fields(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("学生成绩分析系统", level=1)
    document.add_paragraph("课题背景：面向学院成绩统计")
    document.add_paragraph("功能要求：用户登录、成绩录入、统计分析")
    document.add_paragraph("约束条件：本地部署、单用户")
    document.add_paragraph("技术要求")
    document.add_paragraph("Python 数据分析", style="List Bullet")
    document.add_paragraph("论文提纲：摘要、系统设计、结论")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "章节"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "摘要"
    table.cell(1, 1).text = "论文摘要内容"
    document.save(file_path)

    result = parse_brief(file_path)

    assert result.title == "学生成绩分析系统"
    assert "用户登录" in result.requirements
    assert "本地部署" in result.constraints
    assert "Python 数据分析" in result.tech_hints
    assert "摘要" in result.thesis_cues
    assert "学生成绩分析系统" in result.raw_text
    assert result.extraction_snapshot["tables"][0][1][0] == "摘要"
