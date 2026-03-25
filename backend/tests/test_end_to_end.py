from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient

from idea2thesis.config import Settings
from idea2thesis.db import open_connection
from idea2thesis.main import create_app
from idea2thesis.secrets import read_job_secret
from idea2thesis.worker import AsyncJobWorker


def test_create_job_from_uploaded_brief_returns_snapshot_and_artifacts(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        response = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )

    assert response.status_code == 201
    body = response.json()
    assert body["schema_version"] == "v1alpha1"
    assert body["job_id"]
    assert body["status"] == "pending"
    assert body["stage"] == "queued"
    assert body["validation_state"] == "pending"
    assert body["final_disposition"] == "pending"
    assert isinstance(body["artifacts"], list)
    assert body["artifacts"] == []


def test_uploaded_filename_is_sanitized_to_workspace(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        response = client.post(
            "/jobs",
            files={
                "file": (
                    "../../outside.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )

    assert response.status_code == 201
    jobs_dir = settings.jobs_dir
    assert not (tmp_path / "outside.docx").exists()
    stored_files = list(jobs_dir.glob("*/input/*"))
    assert stored_files
    assert all(path.parent.parent.parent == jobs_dir for path in stored_files)
    secret_files = list(settings.secret_dir.glob("*.bin"))
    assert secret_files


def test_job_creation_rejects_unknown_agent_override_role(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        response = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {
                    "rogue": {
                      "use_global": false,
                      "api_key": "rogue-key",
                      "base_url": "https://rogue.example/v1",
                      "model": "gpt-rogue"
                    }
                  }
                }
                """
            },
        )

    assert response.status_code == 422
    assert "unknown agent role" in response.text


def test_job_creation_rejects_missing_effective_api_key(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        response = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )

    assert response.status_code == 422
    assert "missing effective api_key" in response.text


def test_rerun_and_delete_integration(tmp_path: Path) -> None:
    client = TestClient(
        create_app(
            Settings(
                jobs_dir=tmp_path / "jobs",
                api_key="",
                base_url="https://example.com/v1",
                model="gpt-test",
                settings_file=tmp_path / ".idea2thesis" / "settings.json",
                database_path=tmp_path / ".idea2thesis" / "jobs.db",
                secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
                secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
            )
        )
    )

    rerun = client.post("/jobs/source-job/rerun", data={"config": "{}"})
    assert rerun.status_code == 422

    deleted = client.delete("/jobs/source-job")
    assert deleted.status_code == 404


def test_job_creation_does_not_persist_agent_api_keys(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        response = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {
                    "coder": {
                      "use_global": false,
                      "api_key": "coder-secret",
                      "base_url": "https://coder.example.com/v1",
                      "model": "gpt-coder"
                    }
                  }
                }
                """
            },
        )

    assert response.status_code == 201
    job_id = response.json()["job_id"]
    with open_connection(settings) as connection:
        row = connection.execute(
            "SELECT agents_json FROM job_runtime_inputs WHERE job_id = ?",
            (job_id,),
        ).fetchone()
    assert row is not None
    assert "coder-secret" not in str(row[0])
    assert '"api_key"' not in str(row[0])


def test_rerun_writes_fresh_secret_and_uses_submitted_runtime_config(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        created = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {
                    "coder": {
                      "use_global": false,
                      "api_key": "coder-key-1",
                      "base_url": "https://coder.example.com/v1",
                      "model": "gpt-coder-1"
                    }
                  }
                }
                """
            },
        )
    assert created.status_code == 201
    source_job_id = created.json()["job_id"]

    with open_connection(settings) as connection:
        connection.execute(
            "UPDATE jobs SET status = ?, stage = ?, validation_state = ?, final_disposition = ? WHERE id = ?",
            ("completed", "completed", "completed", "completed", source_job_id),
        )
        connection.commit()

    rerun = client.post(
        f"/jobs/{source_job_id}/rerun",
        data={
            "config": """
            {
              "schema_version": "v1alpha1",
              "global": {
                "api_key": "runtime-key-2",
                "base_url": "https://rerun.example.com/v1",
                "model": "gpt-rerun"
              },
              "agents": {
                "coder": {
                  "use_global": false,
                  "api_key": "coder-key-2",
                  "base_url": "https://rerun-coder.example.com/v1",
                  "model": "gpt-rerun-coder"
                }
              }
            }
            """
        },
    )

    assert rerun.status_code == 201
    body = rerun.json()
    assert body["runtime_preset"]["global"]["base_url"] == "https://rerun.example.com/v1"
    assert body["runtime_preset"]["global"]["model"] == "gpt-rerun"
    assert body["runtime_preset"]["agents"]["coder"]["base_url"] == "https://rerun-coder.example.com/v1"
    assert body["runtime_preset"]["agents"]["coder"]["model"] == "gpt-rerun-coder"

    with open_connection(settings) as connection:
        row = connection.execute(
            "SELECT secret_file_path FROM jobs WHERE id = ?",
            (body["job_id"],),
        ).fetchone()
    assert row is not None
    secret_path = Path(str(row[0]))
    assert secret_path.exists()
    envelope = read_job_secret(settings, secret_path)
    assert envelope.global_api_key == "runtime-key-2"
    assert envelope.per_agent_api_keys["coder"] == "coder-key-2"


def test_worker_execution_persists_real_artifacts_and_events(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询、借阅管理")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        created = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )
    assert created.status_code == 201
    job_id = created.json()["job_id"]

    worker = AsyncJobWorker(settings)
    assert worker.run_once() is True

    detail = client.get(f"/jobs/{job_id}")
    events = client.get(f"/jobs/{job_id}/events")
    assert detail.status_code == 200
    assert events.status_code == 200
    body = detail.json()
    event_kinds = [item["kind"] for item in events.json()["items"]]
    assert body["status"] == "completed"
    assert body["final_disposition"] == "completed"
    assert any(item["kind"] == "job_manifest" for item in body["artifacts"])
    assert any(item["kind"] == "code_eval" for item in body["artifacts"])
    assert "verification_started" in event_kinds
    assert "verification_completed" in event_kinds

    manifest_path = tmp_path / "jobs" / job_id / "artifacts" / "final" / "job_manifest.json"
    manifest_text = manifest_path.read_text(encoding="utf-8")
    assert '"final_disposition": "completed"' in manifest_text
    assert "runtime-key" not in manifest_text

    with open_connection(settings) as connection:
        row = connection.execute(
            "SELECT secret_file_path FROM jobs WHERE id = ?",
            (job_id,),
        ).fetchone()
    assert row is not None
    assert row[0] is None


def test_artifact_content_endpoint_returns_registered_text_artifact(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询、借阅管理")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        created = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )
    assert created.status_code == 201
    job_id = created.json()["job_id"]

    worker = AsyncJobWorker(settings)
    assert worker.run_once() is True

    detail = client.get(f"/jobs/{job_id}")
    assert detail.status_code == 200
    preview_artifact = next(
        item for item in detail.json()["artifacts"] if item["kind"] == "project_readme"
    )

    preview = client.get(
        f"/jobs/{job_id}/artifacts/content",
        params={"path": preview_artifact["path"]},
    )

    assert preview.status_code == 200
    body = preview.json()
    assert body["path"] == preview_artifact["path"]
    assert body["content"]
    assert body["truncated"] is False


def test_artifact_content_endpoint_rejects_unregistered_path(tmp_path: Path) -> None:
    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    response = client.get(
        "/jobs/job-1/artifacts/content",
        params={"path": "/tmp/not-registered.txt"},
    )
    assert response.status_code == 404


def test_artifact_download_endpoint_returns_registered_artifact_file(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询、借阅管理")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        created = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )
    assert created.status_code == 201
    job_id = created.json()["job_id"]

    worker = AsyncJobWorker(settings)
    assert worker.run_once() is True

    detail = client.get(f"/jobs/{job_id}")
    assert detail.status_code == 200
    download_artifact = next(
        item for item in detail.json()["artifacts"] if item["kind"] == "project_readme"
    )

    response = client.get(
        f"/jobs/{job_id}/artifacts/download",
        params={"path": download_artifact["path"]},
    )
    assert response.status_code == 200
    assert response.content
    assert response.headers["content-type"].startswith("text/")
    assert "attachment" in response.headers["content-disposition"]


def test_artifact_open_endpoint_executes_system_file_manager_for_registered_artifact(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询、借阅管理")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        created = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )
    assert created.status_code == 201
    job_id = created.json()["job_id"]

    worker = AsyncJobWorker(settings)
    assert worker.run_once() is True

    detail = client.get(f"/jobs/{job_id}")
    assert detail.status_code == 200
    open_artifact = next(
        item for item in detail.json()["artifacts"] if item["kind"] == "project_readme"
    )

    with patch("idea2thesis.services.platform.system", return_value="Darwin"):
        with patch("idea2thesis.services.subprocess.run") as run_mock:
            response = client.post(
                f"/jobs/{job_id}/artifacts/open",
                params={"path": open_artifact["path"]},
            )

    assert response.status_code == 200
    assert response.json()["ok"] is True
    assert response.json()["path"] == open_artifact["path"]
    run_mock.assert_called_once()


def test_artifact_download_endpoint_returns_generated_thesis_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询、借阅管理")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    with file_path.open("rb") as handle:
        created = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )
    assert created.status_code == 201
    job_id = created.json()["job_id"]

    worker = AsyncJobWorker(settings)
    assert worker.run_once() is True

    detail = client.get(f"/jobs/{job_id}")
    assert detail.status_code == 200
    thesis_docx_artifact = next(
        item for item in detail.json()["artifacts"] if item["path"].endswith("thesis_draft.docx")
    )

    response = client.get(
        f"/jobs/{job_id}/artifacts/download",
        params={"path": thesis_docx_artifact["path"]},
    )

    assert response.status_code == 200
    assert response.content.startswith(b"PK")
    assert "thesis_draft.docx" in response.headers["content-disposition"]


def test_persisted_thesis_cover_settings_are_written_into_generated_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "brief.docx"
    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询、借阅管理")
    document.save(file_path)

    settings = Settings(
        jobs_dir=tmp_path / "jobs",
        api_key="",
        base_url="https://example.com/v1",
        model="gpt-test",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )
    client = TestClient(create_app(settings))
    saved = client.put(
        "/settings",
        json={
            "schema_version": "v1alpha1",
            "global": {
                "base_url": "https://example.com/v1",
                "model": "gpt-test",
                "thesis_cover": {
                    "school": "示例大学",
                    "department": "计算机学院",
                    "major": "软件工程",
                    "student_name": "张三",
                    "student_id": "20240001",
                    "advisor": "李老师",
                },
            },
            "agents": {},
        },
    )
    assert saved.status_code == 200

    with file_path.open("rb") as handle:
        created = client.post(
            "/jobs",
            files={
                "file": (
                    "brief.docx",
                    handle.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "config": """
                {
                  "schema_version": "v1alpha1",
                  "global": {
                    "api_key": "runtime-key",
                    "base_url": "https://example.com/v1",
                    "model": "gpt-test"
                  },
                  "agents": {}
                }
                """
            },
        )
    assert created.status_code == 201
    job_id = created.json()["job_id"]

    worker = AsyncJobWorker(settings)
    assert worker.run_once() is True

    generated_docx = (
        tmp_path / "jobs" / job_id / "artifacts" / "agent" / "writer" / "thesis_draft.docx"
    )
    thesis_document = Document(generated_docx)
    paragraphs = "\n".join(
        paragraph.text.strip()
        for paragraph in thesis_document.paragraphs
        if paragraph.text.strip()
    )
    assert "学校：示例大学" in paragraphs
    assert "学院：计算机学院" in paragraphs
    assert "学生姓名：张三" in paragraphs
    assert "学号：20240001" in paragraphs
    assert "指导教师：李老师" in paragraphs
