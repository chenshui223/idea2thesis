from pathlib import Path

from docx import Document

from idea2thesis.config import Settings
from idea2thesis.db import initialize_database
from idea2thesis.job_store import JobStore
from idea2thesis.parser import parse_brief
from idea2thesis.secrets import JobSecretEnvelope, write_job_secret
from idea2thesis.storage import JobStorage
from idea2thesis.worker import AsyncJobWorker


def build_settings(tmp_path: Path) -> Settings:
    return Settings(
        jobs_dir=tmp_path / "jobs",
        database_path=tmp_path / ".idea2thesis" / "jobs.db",
        settings_file=tmp_path / ".idea2thesis" / "settings.json",
        secret_key_path=tmp_path / ".idea2thesis" / "secret.key",
        secret_dir=tmp_path / ".idea2thesis" / "job-secrets",
    )


def seed_pending_job(tmp_path: Path, job_id: str = "job-1") -> tuple[Settings, JobStore]:
    settings = build_settings(tmp_path)
    initialize_database(settings)
    storage = JobStorage(settings.jobs_dir)
    paths = storage.create_job_workspace(job_id)

    document = Document()
    document.add_heading("图书管理系统", level=1)
    document.add_paragraph("功能要求：用户登录、图书查询")
    input_path = paths.input_dir / "brief.docx"
    document.save(input_path)
    brief = parse_brief(input_path)
    (paths.parsed_dir / "brief.json").write_text(brief.model_dump_json(indent=2), encoding="utf-8")

    secret_path = write_job_secret(
        settings,
        job_id,
        JobSecretEnvelope(global_api_key="runtime-key", per_agent_api_keys={}),
    )
    store = JobStore(settings)
    store.create_job(
        job_id=job_id,
        brief_title=brief.title,
        input_file_path=str(input_path),
        workspace_path=str(paths.workspace_dir),
        secret_file_path=str(secret_path),
        runtime_inputs={
            "global_base_url": "https://example.com/v1",
            "global_model": "gpt-test",
            "agents_json": "{}",
            "api_key_required": True,
        },
        agents=["advisor", "coder", "writer", "requirements_reviewer", "engineering_reviewer", "delivery_reviewer", "code_eval", "doc_check"],
    )
    return settings, store


def test_worker_claims_pending_job_and_persists_completion(tmp_path: Path) -> None:
    settings, store = seed_pending_job(tmp_path)

    worker = AsyncJobWorker(settings)
    worker.run_once()

    snapshot = store.get_job("job-1")
    assert snapshot.status == "completed"
    assert any(item.kind == "verification_report" for item in snapshot.artifacts)
    assert not store.get_job_record("job-1").secret_file_path


def test_worker_startup_reconciles_stale_running_jobs_only(tmp_path: Path) -> None:
    settings, store = seed_pending_job(tmp_path, job_id="stale-running")
    worker = AsyncJobWorker(settings)

    store.register_worker_session("worker-stale")
    store.claim_next_job("worker-stale")

    pending_settings, pending_store = seed_pending_job(tmp_path, job_id="still-pending")
    assert pending_settings == settings
    assert pending_store.get_job("still-pending").status == "pending"

    worker.reconcile_startup_state()

    assert store.get_job("stale-running").status == "interrupted"
    assert store.get_job("still-pending").status == "pending"
    assert not store.get_job_record("stale-running").secret_file_path
