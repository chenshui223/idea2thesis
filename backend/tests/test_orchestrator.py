import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from idea2thesis.contracts import ParsedBrief
from idea2thesis.executor import LocalCommandExecutor
from idea2thesis.orchestrator import SupervisorOrchestrator
from idea2thesis.providers.base import CompletionProvider
from idea2thesis.providers.runner import AgentProviderConfig
from idea2thesis.storage import JobPaths, JobStorage


def sample_brief(title: str = "图书管理系统") -> ParsedBrief:
    return ParsedBrief(
        title=title,
        requirements=["用户登录", "图书查询", "借阅管理"],
        constraints=["本地部署"],
        tech_hints=["Python", "FastAPI"],
        thesis_cues=["摘要", "系统设计"],
        raw_text=title,
        extraction_snapshot={"paragraphs": [title]},
    )


def seeded_job_paths(tmp_path: Path, job_id: str) -> JobPaths:
    storage = JobStorage(tmp_path / "jobs")
    return storage.create_job_workspace(job_id)


def artifact_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def artifact_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8")


class FakeProvider:
    def __init__(self, content: str | Exception) -> None:
        self.content = content

    def complete(self, prompt: str) -> str:
        if isinstance(self.content, Exception):
            raise self.content
        return self.content


class PromptRouterProvider:
    def __init__(self, responses: dict[str, str]) -> None:
        self.responses = responses

    def complete(self, prompt: str) -> str:
        for marker, response in self.responses.items():
            if marker in prompt:
                return response
        raise AssertionError(f"unexpected prompt: {prompt}")


def test_supervisor_builds_plan_with_review_criteria_and_tasks() -> None:
    orchestrator = SupervisorOrchestrator()
    brief = ParsedBrief(
        title="学生成绩分析系统",
        requirements=["用户登录", "统计分析"],
        constraints=["本地部署"],
        tech_hints=["Python"],
        thesis_cues=["摘要"],
        raw_text="学生成绩分析系统",
        extraction_snapshot={"paragraphs": ["学生成绩分析系统"]},
    )
    plan = orchestrator.build_plan(brief)
    assert plan.project_category == "data_analysis_project"
    assert plan.stack_policy == "python-data"
    assert plan.review_criteria == [
        "requirements_alignment",
        "engineering_quality",
        "delivery_readiness",
    ]
    assert plan.retries["code_eval"] == 1
    assert [task.role for task in plan.tasks] == [
        "advisor",
        "coder",
        "writer",
        "requirements_reviewer",
        "engineering_reviewer",
        "delivery_reviewer",
        "code_eval",
        "doc_check",
    ]


def test_run_job_persists_real_stage_artifacts_and_manifest(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator()
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-1")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-1",
        brief,
        paths,
        executor,
        thesis_cover={
            "school": "示例大学",
            "department": "计算机学院",
            "major": "软件工程",
            "student_name": "张三",
            "student_id": "20240001",
            "advisor": "李老师",
        },
    )
    assert snapshot.status == "completed"
    assert snapshot.stage == "completed"
    assert artifact_json(paths.artifacts_dir / "agent" / "advisor" / "advisor_plan.json")["agent_role"] == "advisor"
    assert artifact_json(paths.artifacts_dir / "agent" / "coder" / "code_summary.json")["agent_role"] == "coder"
    assert artifact_markdown(paths.artifacts_dir / "agent" / "writer" / "thesis_draft.md").startswith("#")
    thesis_docx = paths.artifacts_dir / "agent" / "writer" / "thesis_draft.docx"
    assert thesis_docx.exists()
    document = Document(thesis_docx)
    paragraphs = "\n".join(
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    )
    assert "图书管理系统" in paragraphs
    assert "本科毕业设计（论文）" in paragraphs
    assert "学生姓名：张三" in paragraphs
    assert "学号：20240001" in paragraphs
    assert "指导教师：李老师" in paragraphs
    assert "目录" in paragraphs
    assert "目录待在 Word 中更新" in paragraphs
    title_paragraph = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "图书管理系统 Thesis Draft"
    )
    assert title_paragraph.paragraph_format.line_spacing_rule is None
    body_paragraph = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("图书管理系统 面向本地单用户毕业设计场景")
    )
    assert body_paragraph.paragraph_format.first_line_indent == Pt(21)
    assert body_paragraph.paragraph_format.line_spacing == 1.5
    assert artifact_json(paths.artifacts_dir / "final" / "job_manifest.json")["final_disposition"] == "completed"


def test_doc_check_blocks_when_required_sections_are_missing(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator()
    brief = ParsedBrief(
        title="图书管理系统",
        requirements=[],
        constraints=["本地部署"],
        tech_hints=["FastAPI"],
        thesis_cues=[],
        raw_text="图书管理系统",
        extraction_snapshot={"paragraphs": ["图书管理系统"]},
    )
    paths = seeded_job_paths(tmp_path, "job-2")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job("job-2", brief, paths, executor)

    doc_check = artifact_json(paths.artifacts_dir / "verification" / "doc_check.json")
    manifest = artifact_json(paths.artifacts_dir / "final" / "job_manifest.json")
    assert doc_check["status"] == "must_fix"
    assert snapshot.status == "blocked"
    assert manifest["final_disposition"] == "blocked"


def test_code_eval_failure_marks_job_failed(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator()
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-3")
    executor = LocalCommandExecutor(paths.workspace_dir)
    (paths.workspace_dir / "fail.py").write_text("raise SystemExit(2)\n", encoding="utf-8")

    snapshot = orchestrator.run_job("job-3", brief, paths, executor)

    code_eval = artifact_json(paths.artifacts_dir / "verification" / "code_eval.json")
    manifest = artifact_json(paths.artifacts_dir / "final" / "job_manifest.json")
    assert code_eval["status"] == "failed"
    assert snapshot.status == "failed"
    assert manifest["final_disposition"] == "failed"


def test_provider_failure_falls_back_to_deterministic_generation(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: FakeProvider(RuntimeError(f"{config.role} down"))
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-4")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-4",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    advisor = artifact_json(paths.artifacts_dir / "agent" / "advisor" / "advisor_plan.json")
    code_summary = artifact_json(paths.artifacts_dir / "agent" / "coder" / "code_summary.json")
    thesis_text = artifact_markdown(paths.artifacts_dir / "agent" / "writer" / "thesis_draft.md")
    assert snapshot.status == "completed"
    assert "fallback" in advisor["summary"].lower()
    assert "fallback" in code_summary["summary"].lower()
    assert thesis_text.startswith("#")


def test_provider_structured_output_populates_advisor_coder_and_writer_artifacts(
    tmp_path: Path,
) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: PromptRouterProvider(
            {
                "毕业设计指导老师 agent": """
                {
                  "summary": "structured advisor summary",
                  "project_summary": "面向课程设计的一键式数据分析项目生成器。",
                  "recommended_stack": "python-data-plus",
                  "module_breakdown": ["data_ingest", "analysis_pipeline", "report_export"],
                  "implementation_priorities": ["解析设计书", "生成代码", "执行评测"],
                  "writing_priorities": ["摘要", "方案设计", "实验结果"],
                  "risks": ["训练数据不足"],
                  "coder_directives": ["生成可运行脚本", "输出评测命令"],
                  "writer_directives": ["围绕实验结果展开论文初稿"]
                }
                """,
                "你是 coder agent": """
                ```json
                {
                  "summary": "structured coder summary",
                  "generated_files": ["README.md", "src/pipeline.py", "docs/usage.md"],
                  "chosen_stack": "python-data-plus",
                  "run_commands": ["python -m app.cli run"],
                  "test_commands": ["pytest -q", "python -m app.cli smoke"],
                  "known_limitations": ["需要用户补充真实数据集"]
                }
                ```
                """,
                "你是 writer agent": """
                {
                  "summary": "structured writer summary",
                  "title": "图书管理系统论文初稿",
                  "sections": ["摘要", "需求分析", "系统设计", "实现概述", "测试与验证", "结论"],
                  "abstract": "本文提出一个支持本地部署的数据分析类毕业设计生成系统。",
                  "requirements_analysis": "需求覆盖文档解析、代码生成与论文生成。",
                  "system_design": "系统采用多 agent 协同架构。",
                  "implementation_overview": "实现阶段集成 provider 调用与回退策略。",
                  "testing_validation": "通过本地命令对关键流程进行验证。",
                  "conclusion": "结果表明该方案适合作为毕业设计初稿基座。",
                  "design_report": {
                    "summary": "structured design report summary",
                    "goal": "沉淀可复用的一键生成流程。",
                    "module_breakdown": ["advisor", "coder", "writer"],
                    "delivery_notes": "交付包含代码、文档与验证结果。"
                  }
                }
                """,
            }
        )
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-5")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-5",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    advisor = artifact_json(paths.artifacts_dir / "agent" / "advisor" / "advisor_plan.json")
    code_summary = artifact_json(paths.artifacts_dir / "agent" / "coder" / "code_summary.json")
    thesis_text = artifact_markdown(paths.artifacts_dir / "agent" / "writer" / "thesis_draft.md")
    design_text = artifact_markdown(paths.artifacts_dir / "agent" / "writer" / "design_report.md")

    assert snapshot.status == "completed"
    assert advisor["summary"] == "structured advisor summary"
    assert advisor["recommended_stack"] == "python-data-plus"
    assert advisor["module_breakdown"] == ["data_ingest", "analysis_pipeline", "report_export"]
    assert code_summary["summary"] == "structured coder summary"
    assert code_summary["generated_files"] == ["README.md", "src/pipeline.py", "docs/usage.md"]
    assert code_summary["run_commands"] == ["python -m app.cli run"]
    assert "本文提出一个支持本地部署的数据分析类毕业设计生成系统。" in thesis_text
    assert "系统采用多 agent 协同架构。" in thesis_text
    assert "交付包含代码、文档与验证结果。" in design_text


def test_provider_malformed_structured_output_falls_back_to_deterministic_fields(
    tmp_path: Path,
) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: PromptRouterProvider(
            {
                "毕业设计指导老师 agent": "{not-valid-json}",
                "你是 coder agent": "plain summary only",
                "你是 writer agent": "```json\n{\"summary\": \"writer only\"}\n```",
            }
        )
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-6")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-6",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    advisor = artifact_json(paths.artifacts_dir / "agent" / "advisor" / "advisor_plan.json")
    code_summary = artifact_json(paths.artifacts_dir / "agent" / "coder" / "code_summary.json")
    thesis_text = artifact_markdown(paths.artifacts_dir / "agent" / "writer" / "thesis_draft.md")

    assert snapshot.status == "completed"
    assert advisor["recommended_stack"] == "fastapi-react"
    assert advisor["module_breakdown"] == [
        "authentication",
        "catalog_search",
        "workflow_management",
        "reporting",
    ]
    assert code_summary["generated_files"] == ["README.md"]
    assert "图书管理系统 面向本地单用户毕业设计场景" in thesis_text


def test_provider_coder_payload_writes_workspace_files_and_plan(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: PromptRouterProvider(
            {
                "毕业设计指导老师 agent": """
                {
                  "summary": "advisor ok",
                  "recommended_stack": "python-data-plus"
                }
                """,
                "你是 coder agent": """
                {
                  "summary": "coder wrote files",
                  "generated_files": ["README.md", "src/pipeline.py", "docs/usage.md"],
                  "chosen_stack": "python-data-plus",
                  "run_commands": ["python src/pipeline.py"],
                  "test_commands": ["pytest -q"],
                  "known_limitations": ["需要补充真实样本"],
                  "implementation_plan_markdown": "# Custom Plan\\n\\n## Steps\\n- ingest\\n- analyze\\n",
                  "workspace_files": [
                    {
                      "path": "README.md",
                      "content": "# Provider Workspace\\n\\nRun `python src/pipeline.py`\\n"
                    },
                    {
                      "path": "src/pipeline.py",
                      "content": "def main():\\n    print('provider generated')\\n\\n\\nif __name__ == '__main__':\\n    main()\\n"
                    },
                    {
                      "path": "docs/usage.md",
                      "content": "# Usage\\n\\nExecute the generated pipeline locally.\\n"
                    }
                  ]
                }
                """,
                "你是 writer agent": """
                {
                  "summary": "writer ok",
                  "abstract": "生成文档内容。"
                }
                """,
            }
        )
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-7")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-7",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    code_summary = artifact_json(paths.artifacts_dir / "agent" / "coder" / "code_summary.json")
    implementation_plan = artifact_markdown(paths.artifacts_dir / "agent" / "coder" / "implementation_plan.md")
    readme_text = artifact_markdown(paths.workspace_dir / "README.md")
    pipeline_text = artifact_markdown(paths.workspace_dir / "src" / "pipeline.py")
    usage_text = artifact_markdown(paths.workspace_dir / "docs" / "usage.md")

    assert snapshot.status == "completed"
    assert code_summary["summary"] == "coder wrote files"
    assert code_summary["generated_files"] == ["README.md", "src/pipeline.py", "docs/usage.md"]
    assert implementation_plan.startswith("# Custom Plan")
    assert readme_text.startswith("# Provider Workspace")
    assert "provider generated" in pipeline_text
    assert "Execute the generated pipeline locally." in usage_text


def test_provider_workspace_file_generation_blocks_path_escape(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: PromptRouterProvider(
            {
                "毕业设计指导老师 agent": """
                {
                  "summary": "advisor ok"
                }
                """,
                "你是 coder agent": """
                {
                  "summary": "coder attempted escape",
                  "generated_files": ["README.md"],
                  "workspace_files": [
                    {
                      "path": "../outside.txt",
                      "content": "escaped"
                    }
                  ]
                }
                """,
                "你是 writer agent": """
                {
                  "summary": "writer ok",
                  "abstract": "生成文档内容。"
                }
                """,
            }
        )
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-8")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-8",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    code_summary = artifact_json(paths.artifacts_dir / "agent" / "coder" / "code_summary.json")

    assert snapshot.status == "completed"
    assert not (paths.root_dir / "outside.txt").exists()
    assert code_summary["generated_files"] == ["README.md"]
    assert "deterministic scaffold" in code_summary["known_limitations"][0]


def test_provider_writer_payload_writes_supporting_documents(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: PromptRouterProvider(
            {
                "毕业设计指导老师 agent": """
                {
                  "summary": "advisor ok"
                }
                """,
                "你是 coder agent": """
                {
                  "summary": "coder ok"
                }
                """,
                "你是 writer agent": """
                {
                  "summary": "writer published docs",
                  "title": "图书管理系统论文初稿",
                  "abstract": "本文围绕本地单用户毕业设计生成流程展开。",
                  "requirements_analysis": "需求覆盖文档解析、代码生成、论文输出。",
                  "system_design": "系统由 advisor、coder、writer 和 reviewer 协同完成。",
                  "implementation_overview": "通过 provider 结果直接落盘到 workspace。",
                  "testing_validation": "采用本地命令与文档校验联合验证。",
                  "conclusion": "该流程可作为毕业设计初稿生产基座。",
                  "workspace_files": [
                    {
                      "path": "docs/答辩提纲.md",
                      "content": "# 答辩提纲\\n\\n- 项目背景\\n- 系统设计\\n- 实验结果\\n"
                    },
                    {
                      "path": "docs/交付说明.md",
                      "content": "# 交付说明\\n\\n包含代码、论文初稿与验证结果。\\n"
                    }
                  ],
                  "design_report": {
                    "summary": "design ok",
                    "goal": "固化文档交付模板。",
                    "module_breakdown": ["advisor", "coder", "writer"],
                    "delivery_notes": "交付需附带答辩提纲与交付说明。"
                  }
                }
                """,
            }
        )
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-9")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-9",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    thesis_text = artifact_markdown(paths.artifacts_dir / "agent" / "writer" / "thesis_draft.md")
    design_text = artifact_markdown(paths.artifacts_dir / "agent" / "writer" / "design_report.md")
    defense_outline = artifact_markdown(paths.workspace_dir / "docs" / "答辩提纲.md")
    delivery_notes = artifact_markdown(paths.workspace_dir / "docs" / "交付说明.md")

    assert snapshot.status == "completed"
    assert "本文围绕本地单用户毕业设计生成流程展开。" in thesis_text
    assert "交付需附带答辩提纲与交付说明。" in design_text
    assert defense_outline.startswith("# 答辩提纲")
    assert "包含代码、论文初稿与验证结果。" in delivery_notes


def test_provider_writer_workspace_file_generation_blocks_path_escape(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: PromptRouterProvider(
            {
                "毕业设计指导老师 agent": """
                {
                  "summary": "advisor ok"
                }
                """,
                "你是 coder agent": """
                {
                  "summary": "coder ok"
                }
                """,
                "你是 writer agent": """
                {
                  "summary": "writer attempted escape",
                  "abstract": "论文摘要。",
                  "workspace_files": [
                    {
                      "path": "../writer-outside.md",
                      "content": "escaped"
                    }
                  ]
                }
                """,
            }
        )
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-10")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-10",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    assert snapshot.status == "completed"
    assert not (paths.root_dir / "writer-outside.md").exists()


def test_provider_generated_workspace_files_are_listed_as_artifacts(tmp_path: Path) -> None:
    orchestrator = SupervisorOrchestrator(
        provider_factory=lambda config: PromptRouterProvider(
            {
                "毕业设计指导老师 agent": """
                {
                  "summary": "advisor ok"
                }
                """,
                "你是 coder agent": """
                {
                  "summary": "coder wrote files",
                  "generated_files": ["README.md", "src/pipeline.py"],
                  "workspace_files": [
                    {
                      "path": "README.md",
                      "content": "# Provider Workspace\\n"
                    },
                    {
                      "path": "src/pipeline.py",
                      "content": "print('provider generated')\\n"
                    }
                  ]
                }
                """,
                "你是 writer agent": """
                {
                  "summary": "writer published docs",
                  "abstract": "论文摘要。",
                  "workspace_files": [
                    {
                      "path": "docs/答辩提纲.md",
                      "content": "# 答辩提纲\\n"
                    }
                  ]
                }
                """,
            }
        )
    )
    brief = sample_brief(title="图书管理系统")
    paths = seeded_job_paths(tmp_path, "job-11")
    executor = LocalCommandExecutor(paths.workspace_dir)

    snapshot = orchestrator.run_job(
        "job-11",
        brief,
        paths,
        executor,
        provider_configs={
            "advisor": AgentProviderConfig(
                role="advisor",
                api_key="advisor-key",
                base_url="https://example.com/v1",
                model="gpt-advisor",
            ),
            "coder": AgentProviderConfig(
                role="coder",
                api_key="coder-key",
                base_url="https://example.com/v1",
                model="gpt-coder",
            ),
            "writer": AgentProviderConfig(
                role="writer",
                api_key="writer-key",
                base_url="https://example.com/v1",
                model="gpt-writer",
            ),
        },
    )

    workspace_artifacts = [item for item in snapshot.artifacts if item.kind == "workspace_file"]
    workspace_paths = sorted(item.path for item in workspace_artifacts)

    assert snapshot.status == "completed"
    assert workspace_paths == sorted(
        [
            str(paths.workspace_dir / "README.md"),
            str(paths.workspace_dir / "src" / "pipeline.py"),
            str(paths.workspace_dir / "docs" / "答辩提纲.md"),
        ]
    )
